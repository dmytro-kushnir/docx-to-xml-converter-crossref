from docx import Document
import os
import lxml.etree as etree

NAMESPACE = {'ns': 'http://www.crossref.org/schema/4.4.2'}

def parse_xml(xml_path):
    """Parse XML file and return its root element."""
    tree = etree.parse(xml_path)
    return tree.getroot()

def parse_journal_metadata(xml_data):
    """Extract journal title, volume, issue, year, ISSN, and URLs from XML."""
    journal_metadata = xml_data.xpath('//ns:journal_metadata', namespaces=NAMESPACE)[0]
    journal_title = journal_metadata.find('ns:full_title', namespaces=NAMESPACE).text

    issn_node = journal_metadata.find('ns:issn[@media_type="print"]', namespaces=NAMESPACE)
    issn = issn_node.text if issn_node is not None else "N/A"

    journal_url_node = journal_metadata.find('ns:doi_data/ns:resource', namespaces=NAMESPACE)
    journal_url = journal_url_node.text if journal_url_node is not None else "N/A"

    issue_metadata = xml_data.xpath('//ns:journal_issue', namespaces=NAMESPACE)[0]
    volume = issue_metadata.find('ns:journal_volume/ns:volume', namespaces=NAMESPACE).text
    issue = issue_metadata.find('ns:issue', namespaces=NAMESPACE).text

    issue_url_node = issue_metadata.find('ns:doi_data/ns:resource', namespaces=NAMESPACE)
    issue_url = issue_url_node.text if issue_url_node is not None else "N/A"

    # Extract the year
    year_node = issue_metadata.find('ns:publication_date/ns:year', namespaces=NAMESPACE)
    year = year_node.text if year_node is not None else "N/A"

    return journal_title, volume, issue, year, issn, journal_url, issue_url

def parse_articles(xml_data):
    """Extract articles with their authors, titles, pages, DOIs, and URLs from XML."""
    articles = xml_data.xpath('//ns:journal_article', namespaces=NAMESPACE)
    articles_data = []

    for i, article in enumerate(articles, start=1):
        title_node = article.find('ns:titles/ns:title', namespaces=NAMESPACE)
        title_text = title_node.text if title_node is not None else "N/A"

        authors = ", ".join(
            f"{author.find('ns:given_name', namespaces=NAMESPACE).text or ''} {author.find('ns:surname', namespaces=NAMESPACE).text or ''}".strip()
            for author in article.findall('ns:contributors/ns:person_name', namespaces=NAMESPACE)
        ) or "N/A"

        # Extract pages
        pages_node = article.find('ns:pages', namespaces=NAMESPACE)
        first_page = pages_node.find('ns:first_page', namespaces=NAMESPACE).text if pages_node is not None and pages_node.find('ns:first_page', namespaces=NAMESPACE) is not None else "N/A"
        last_page = pages_node.find('ns:last_page', namespaces=NAMESPACE).text if pages_node is not None and pages_node.find('ns:last_page', namespaces=NAMESPACE) is not None else "N/A"
        page_range = f"{first_page}-{last_page}" if first_page != "N/A" and last_page != "N/A" else "N/A"

        # Extract article DOI and URL
        url_node = article.find('ns:doi_data/ns:resource', namespaces=NAMESPACE)
        url_text = url_node.text if url_node is not None else "N/A"

        doi_node = article.find('ns:doi_data/ns:doi', namespaces=NAMESPACE)
        doi_text = doi_node.text if doi_node is not None else "N/A"

        articles_data.append((i, authors, title_text, page_range, url_text, doi_text))

    return articles_data

def create_contents_docx(xml_name, output_path):
    """Generate contents_eng.docx with journal metadata and article listing."""
    xml_root = parse_xml(xml_name)
    journal_title, volume, issue, year, _, _, _ = parse_journal_metadata(xml_root)
    articles_data = parse_articles(xml_root)

    doc = Document()
    doc.add_paragraph("ЗМІСТ", style="Title")
    doc.add_paragraph(
        f"наукового журналу”\n"
        f"“{journal_title}”\n"
        f"Вип. {volume}, №{issue}, {year} рік", style="Normal"
    )

    doc.add_paragraph("\n")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Автори"
    hdr_cells[2].text = "Назви статей"

    for article in articles_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(article[0])
        row_cells[1].text = article[1]
        row_cells[2].text = article[2]

    doc.save(output_path)
    print(f"Contents document saved to {output_path}")


def create_doi_letter_docx(xml_name, output_path):
    """Generate DOI letter document with journal metadata and articles."""
    xml_root = parse_xml(xml_name)
    journal_title, volume, issue, year, issn, journal_url, issue_url = parse_journal_metadata(xml_root)
    articles_data = parse_articles(xml_root)

    doc = Document()
    doc.add_heading(journal_title, level=1)
    doc.add_paragraph(f"ISSN: {issn}")
    doc.add_paragraph(f"Volume: {volume}, Issue: {issue}, Year: {year}")
    doc.add_paragraph(f"URL: {journal_url}")
    doc.add_paragraph(f"Contents URL: {issue_url}")

    # Add a table for articles
    doc.add_paragraph("\nArticles:\n")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Define table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Authors, Title, and URL"
    hdr_cells[2].text = "Pages"
    hdr_cells[3].text = "DOI"

    for article in articles_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(article[0])  # Article number
        row_cells[1].text = f"{article[1]}. {article[2]}\n{article[4]}"  # Authors, title, and URL
        row_cells[2].text = article[3]  # Pages
        row_cells[3].text = article[5]  # DOI

    doc.save(output_path)
    print(f"DOI Letter document saved to {output_path}")

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(current_dir)
    output_dir = os.path.join(root_dir, "output")

    doi_output = os.path.join(output_dir, "doi_letter.docx")
    contents_output = os.path.join(output_dir, "contents_eng.docx")
    xml_output_name = os.path.join(output_dir, "crossref.xml")

    # Generate documents
    create_doi_letter_docx(xml_output_name, doi_output)
    create_contents_docx(xml_output_name, contents_output)
