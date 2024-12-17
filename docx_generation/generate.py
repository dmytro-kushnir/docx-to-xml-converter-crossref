from docx import Document

def create_docx_with_xml(xml_path, output_path):
    """
    Create a new DOCX document based on XML data.

    :param xml_path: Path to the XML file containing data.
    :param output_path: Path to save the created DOCX file.
    """
    import lxml.etree as etree

    # Parse the XML file
    tree = etree.parse(xml_path)
    namespace = {'ns': 'http://www.crossref.org/schema/4.4.2'}

    # Extract journal metadata
    journal_metadata = tree.xpath('//ns:journal_metadata', namespaces=namespace)[0]
    journal_title = journal_metadata.find('ns:full_title', namespaces=namespace).text
    issn = journal_metadata.find('ns:issn[@media_type="print"]', namespaces=namespace).text
    journal_url = journal_metadata.find('ns:doi_data/ns:resource', namespaces=namespace).text

    # Extract issue metadata
    issue_metadata = tree.xpath('//ns:journal_issue', namespaces=namespace)[0]
    volume = issue_metadata.find('ns:journal_volume/ns:volume', namespaces=namespace).text
    issue = issue_metadata.find('ns:issue', namespaces=namespace).text
    issue_url = issue_metadata.find('ns:doi_data/ns:resource', namespaces=namespace).text

    # Extract articles
    articles = tree.xpath('//ns:journal_article', namespaces=namespace)
    articles_data = []
    for article in articles:
        title = article.find('ns:titles/ns:title', namespaces=namespace)
        title_text = title.text if title is not None else "N/A"

        authors = ", ".join(
            f"{author.find('ns:given_name', namespaces=namespace).text or ''} {author.find('ns:surname', namespaces=namespace).text or ''}".strip()
            for author in article.findall('ns:contributors/ns:person_name', namespaces=namespace)
        ) or "N/A"

        pages = article.find('ns:pages', namespaces=namespace)
        first_page = pages.find('ns:first_page', namespaces=namespace).text if pages is not None and pages.find('ns:first_page', namespaces=namespace) is not None else "N/A"
        last_page = pages.find('ns:last_page', namespaces=namespace).text if pages is not None and pages.find('ns:last_page', namespaces=namespace) is not None else "N/A"
        page_range = f"{first_page}-{last_page}" if first_page != "N/A" and last_page != "N/A" else "N/A"

        url = article.find('ns:doi_data/ns:resource', namespaces=namespace)
        url_text = url.text if url is not None else "N/A"

        doi = article.find('ns:doi_data/ns:doi', namespaces=namespace)
        doi_text = doi.text if doi is not None else "N/A"

        articles_data.append((title_text, authors, page_range, url_text, doi_text))

    # Create the new DOCX document
    doc = Document()

    # Add journal metadata
    doc.add_heading(journal_title, level=1)
    doc.add_paragraph(f"ISSN: {issn}")
    doc.add_paragraph(f"Volume: {volume}, Issue: {issue}, Year: 2024")
    doc.add_paragraph(f"URL: {journal_url}")
    doc.add_paragraph(f"Contents URL: {issue_url}")

    # Add a table for articles
    doc.add_paragraph("\nArticles:\n")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Define table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Authors, Title, and URL"
    hdr_cells[2].text = "Pages"
    hdr_cells[3].text = "DOI"

    # Populate the table with articles
    for i, article in enumerate(articles_data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)  # Article number
        row_cells[1].text = f"{article[1]}. {article[0]}\n{article[3]}"  # Authors, title, and URL
        row_cells[2].text = article[2]  # Pages
        row_cells[3].text = article[4]  # DOI

    # Save the document
    doc.save(output_path)
    print(f"Document saved to {output_path}")
