import argparse
import os
import re
import tempfile

from docx import Document

from docx_processing.extractors import extract_authors, extract_english_title, extract_ukrainian_title, extract_literature
from docx_processing.parse import process_multiple_docs
from pdf_generation.generate_pdf import _check_libreoffice_installed, _convert_docx_to_pdf


EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
EMAIL_KEYWORDS_RE = re.compile(
    r"\b(e-?mail|emails?|authors'?|authors’)\b", re.IGNORECASE
)
AFFILIATION_KEYWORDS_RE = re.compile(
    r"\b(university|institute|department|faculty|academy|college|laboratory|centre|center)\b"
    r"|\b(університет|інститут|кафедра|факультет|академія|коледж|лабораторія|центр)\b",
    re.IGNORECASE,
)
NAME_LIKE_RE = re.compile(r"^[A-ZА-ЯІЇЄҐ][A-Za-zА-Яа-яІЇЄҐіїєґ.'’\-\s]+$")
ORCID_RE = re.compile(r"\bORCID\b|\b\d{4}-\d{4}-\d{4}-\d{4}\b", re.IGNORECASE)


def _build_anonymous_line(original_text):
    year_match = re.search(r"(\d{4})\s*$", original_text)
    year = year_match.group(1) if year_match else ""
    if year:
        return f"© Anonymous {year}"
    return "© Anonymous"


def _extract_author_tokens(authors_text):
    if authors_text == "Authors not found.":
        return []
    normalized = re.sub(r"\s+\d{4}$", "", authors_text).strip()
    parts = re.split(r"\s*(?:,|;| and | та |&)\s*", normalized)
    return [part.strip() for part in parts if len(part.strip()) > 2]


def _extract_surnames(authors_text):
    if not authors_text or authors_text == "Authors not found.":
        return []
    surnames = []
    for part in authors_text.split(","):
        tokens = [t for t in part.strip().split() if t]
        if not tokens:
            continue
        candidate = tokens[-1]
        if re.match(r"^[A-Za-zА-Яа-яІЇЄҐіїєґ'\-]{3,}$", candidate):
            surnames.append(candidate.lower())
    return surnames


def _looks_like_email_line(text):
    return bool(EMAIL_RE.search(text) or EMAIL_KEYWORDS_RE.search(text))


def _looks_like_affiliation_line(text):
    return bool(AFFILIATION_KEYWORDS_RE.search(text))


def _looks_like_orcid_line(text):
    return bool(ORCID_RE.search(text))


def _looks_like_name_line(text):
    if EMAIL_RE.search(text):
        return False
    if re.search(r"\d", text):
        return False
    if len(text) > 80:
        return False
    return bool(NAME_LIKE_RE.match(text.strip()))


def _looks_like_initials_name_line(text):
    if re.search(r"\d", text):
        return False
    if len(text) > 120:
        return False
    # Patterns like "Vitiv N. A." or "Н. А. Вітів"
    has_initials = re.search(r"\b[А-ЯA-Z]\.\s*[А-ЯA-Z]\.", text)
    has_word = re.search(r"\b[А-ЯA-Z][A-Za-zА-Яа-яІЇЄҐіїєґ']+\b", text)
    return bool(has_initials and has_word)

def _contains_author_token(text, author_tokens_lower):
    if not author_tokens_lower:
        return False
    lower = text.lower()
    return any(token in lower for token in author_tokens_lower)

def _clear_container(container):
    for para in container.paragraphs:
        if para.text.strip():
            para.text = ""
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = ""


def anonymize_docx(docx_path, output_docx_path):
    doc = Document(docx_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    authors_en = extract_authors(paragraphs, is_ukrainian=False)
    authors_uk = extract_authors(paragraphs, is_ukrainian=True)
    literature_refs = extract_literature(paragraphs)
    title_uk = extract_ukrainian_title(paragraphs)
    title_en = extract_english_title(paragraphs, literature_refs)
    author_tokens = _extract_author_tokens(authors_en) + _extract_author_tokens(authors_uk)
    author_tokens_lower = [token.lower() for token in author_tokens]
    author_tokens_exact = set(author_tokens_lower)
    surnames_lower = _extract_surnames(authors_en) + _extract_surnames(authors_uk)

    para_texts = [para.text.strip() for para in doc.paragraphs]
    copyright_indices = [
        idx for idx, text in enumerate(para_texts) if text.startswith("©")
    ]
    to_clear = set()

    non_empty_indices = [idx for idx, text in enumerate(para_texts) if text]
    top_indices = set(non_empty_indices[:25])
    text_index_map = {}
    for idx, text in enumerate(para_texts):
        if text:
            key = text.strip().lower()
            text_index_map.setdefault(key, []).append(idx)

    for idx, text in enumerate(para_texts):
        if not text:
            continue

        if (
            _looks_like_email_line(text)
            or _looks_like_affiliation_line(text)
            or _looks_like_orcid_line(text)
        ):
            to_clear.add(idx)

        # Any initials-style author line should be cleared globally
        if _looks_like_initials_name_line(text):
            to_clear.add(idx)
            continue

        if len(text) <= 160 and (
            _contains_author_token(text, author_tokens_lower)
            or any(s in text.lower() for s in surnames_lower)
        ):
            to_clear.add(idx)
            continue

        if idx in top_indices and (
            text.lower() in author_tokens_exact
            or any(s in text.lower() for s in surnames_lower)
        ):
            to_clear.add(idx)

    for title in (title_uk, title_en):
        if not title or title.endswith("not found."):
            continue
        title_key = title.strip().lower()
        for idx in text_index_map.get(title_key, []):
            # Clear the next non-empty line if it looks like a name
            for j in range(idx + 1, len(para_texts)):
                next_text = para_texts[j].strip()
                if not next_text:
                    continue
                if _looks_like_name_line(next_text) or _contains_author_token(
                    next_text, author_tokens_lower
                ):
                    to_clear.add(j)
                break

    for idx in copyright_indices:
        for offset in range(1, 5):
            prev_idx = idx - offset
            if prev_idx < 0:
                break
            prev_text = para_texts[prev_idx]
            if not prev_text:
                continue
            if (
                _looks_like_email_line(prev_text)
                or _looks_like_affiliation_line(prev_text)
                or _contains_author_token(prev_text, author_tokens_lower)
            ):
                to_clear.add(prev_idx)

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("©"):
            para.text = _build_anonymous_line(text)
            continue
        if idx in to_clear:
            para.text = ""

    for section in doc.sections:
        _clear_container(section.header)
        _clear_container(section.first_page_header)
        _clear_container(section.even_page_header)
        _clear_container(section.footer)
        _clear_container(section.first_page_footer)
        _clear_container(section.even_page_footer)

    doc.save(output_docx_path)


def process_and_convert_folder(input_folder, output_folder):
    _check_libreoffice_installed()
    os.makedirs(output_folder, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir:
        for index, (filename, _, _, _) in enumerate(
            process_multiple_docs(input_folder), start=1
        ):
            input_path = os.path.join(input_folder, filename)
            temp_docx_path = os.path.join(temp_dir, filename)
            anon_pdf_name = f"anonymous_{index:03d}.pdf"
            anonymize_docx(input_path, temp_docx_path)

            # Convert using LibreOffice: output name is based on DOCX basename
            dummy_output_path = os.path.join(output_folder, filename.replace(".docx", ".pdf"))
            _convert_docx_to_pdf(temp_docx_path, dummy_output_path)

            generated_pdf = os.path.join(
                output_folder,
                os.path.splitext(os.path.basename(temp_docx_path))[0] + ".pdf",
            )
            final_pdf = os.path.join(output_folder, anon_pdf_name)
            if os.path.exists(generated_pdf):
                os.replace(generated_pdf, final_pdf)


def _parse_args():
    parser = argparse.ArgumentParser(
        description="Anonymize DOCX author lines and convert to PDF.",
    )
    parser.add_argument(
        "input_folder",
        nargs="?",
        help="Folder containing DOCX files to anonymize and convert.",
    )
    parser.add_argument(
        "--output-name",
        default="anonymized_pdfs",
        help="Folder name under ./output for generated PDFs.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = _parse_args()
    if args.input_folder:
        input_folder = args.input_folder
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        input_folder = os.path.join(script_dir, "input")

    if not input_folder:
        raise SystemExit("Input folder is required.")

    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_folder = os.path.join(script_dir, "output", args.output_name)

    process_and_convert_folder(input_folder, output_folder)
